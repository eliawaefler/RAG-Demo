"""
ACHTUNG OPENAI STATT MISTRAL
"""

from pandas.plotting import table
import matplotlib.pyplot as plt
import docx
from docx2txt import process as docx_process
from local_LLM import mistral_complete
import camelot
from camelot.core import *
import fitz  # PyMuPDF
from pdf2image import convert_from_path
from PIL import Image
import io
from uuid import uuid4
import os
import json
import pandas as pd
from typing import Any, Dict, List
from docx2pdf import convert
from openai_handler import *

def process_table(df: pd.DataFrame) -> str:
    """Extract important information from a DataFrame using the Mistral model."""

    # Convert DataFrame to CSV string
    table_csv = df.to_csv(index=False)

    # Create a prompt for the Mistral model
    prompt = (
        "Analyze the following table and provide a summary of the key information:\n\n"
        f"{table_csv}\n\n"
        "The summary should include important insights, trends, and any notable data points."
    )

    # Send the prompt to the Mistral model
    #response = mistral_complete(prompt)
    response = gpt4_new(prompt)
    """
    except Exception as e:
        print(f"Error processing table: {e}")
        return "Error processing table"
    """
    return response


def chunk_text(text: str, chunk_size: int = 1000) -> List[str]:
    """Split text into chunks of specified size."""
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]


def save_image(img: Image, path: str) -> None:
    """Save an image to the specified path."""
    img.save(path, format='PNG')


def process_pdf(filepath: str, output_dir: str) -> Dict[str, Any]:
    """Extract information from a PDF file and return it as a dictionary."""

    doc = fitz.open(filepath)
    filename = os.path.basename(filepath)
    json_data = {
        "name": filename,
        "unique_id": str(uuid4()),
        "disziplin": "",
        "doctype": "",
        "hauptkategorie": "",
        "kategorie": "",
        "subkategorie": "",
        "data": []
    }

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)

        """
        # Extract tables
        tables = camelot.read_pdf(filepath, pages=str(page_num + 1))
        for table in tables:
            df = table.df
            table_info = process_table(df)
            json_data["data"].append({
                "id": len(json_data["data"]) + 1,
                "text": table_info,
                "vektor": []
            })
        """

        # Extract images
        images = page.get_images(full=True)
        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            img = Image.open(io.BytesIO(image_bytes))
            img_path = os.path.join(output_dir, f'{filename}_page{page_num}_img{img_index}.png')
            save_image(img, img_path)
            json_data["data"].append({
                "id": len(json_data["data"]) + 1,
                "titel": f'Image from {filename} page {page_num}',
                "text": "",
                "vektor": [],
                "path": img_path
            })

        # Save a screenshot of the page
        page_img = convert_from_path(filepath, first_page=page_num + 1, last_page=page_num + 1)[0]
        screenshot_path = os.path.join(output_dir, f'{filename}_page{page_num}.png')
        save_image(page_img, screenshot_path)
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "titel": f'Screenshot of page {page_num}',
            "text": "",
            "vektor": [],
            "path": screenshot_path
        })

        # Extract text
        text = page.get_text()
        chunks = chunk_text(text)
        for idx, chunk in enumerate(chunks):
            json_data["data"].append({
                "id": len(json_data["data"]) + 1,
                "text": chunk,
                "vektor": []
            })
    """    
    except Exception as e:
    print(f"Error processing PDF {filepath}: {e}")
    return {}
    """
    return json_data


def process_docx(filepath: str, output_dir: str) -> Dict[str, Any]:
    """Extract information from a DOCX file and return it as a dictionary."""
    doc = docx.Document(filepath)
    filename = os.path.basename(filepath)
    json_data = {
        "name": filename,
        "unique_id": str(uuid4()),
        "disziplin": "",
        "doctype": "",
        "hauptkategorie": "",
        "kategorie": "",
        "subkategorie": "",
        "data": []
    }

    # Extract tables
    for table in doc.tables:
        data = []
        for row in table.rows:
            text = [cell.text for cell in row.cells]
            data.append(text)
        df = pd.DataFrame(data)
        table_info = process_table(df)
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "text": table_info,
            "vektor": []
        })

    """
    # for the screenshot
    # Convert the DOCX file to PDF
    pdf_path = os.path.join(output_dir, f'/temporary_pdfs/{filename}.pdf')
    convert(filepath, pdf_path)

    # Convert each page of the PDF to an image
    pdf_doc = fitz.open(pdf_path)
    for page_num in range(pdf_doc.page_count):
        page_img = convert_from_path(pdf_path, first_page=page_num + 1, last_page=page_num + 1)[0]
        screenshot_path = os.path.join(output_dir, f'{filename}_page{page_num}.png')
        save_image(page_img, screenshot_path)
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "titel": f'Screenshot of page {page_num}',
            "text": "",
            "vektor": [],
            "path": screenshot_path
        })
    """

    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = docx_process(filepath, os.path.join(output_dir, f'{filename}_{rel.target_ref}'))
            img_path = os.path.join(output_dir, f'{filename}_{rel.target_ref}.png')
            save_image(img, img_path)
            json_data["data"].append({
                "id": len(json_data["data"]) + 1,
                "titel": f'Image from {filename}',
                "text": "",
                "vektor": [],
                "path": img_path
            })

    # Extract text
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)

    chunks = chunk_text("\n".join(full_text))
    for idx, chunk in enumerate(chunks):
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "text": chunk,
            "vektor": []
        })

    """
    except Exception as e:
    print(f"Error processing DOCX {filepath}: {e}")
    return {}
    """

    return json_data


def process_xlsx(filepath: str, output_dir: str) -> Dict[str, Any]:
    """Extract information from an XLSX file and return it as a dictionary."""
    filename = os.path.basename(filepath)
    json_data = {
        "name": filename,
        "unique_id": str(uuid4()),
        "disziplin": "",
        "doctype": "",
        "hauptkategorie": "",
        "kategorie": "",
        "subkategorie": "",
        "data": []
    }
    xls = pd.ExcelFile(filepath)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(filepath, sheet_name=sheet_name)
        # Extract table information
        table_info = process_table(df)
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "text": table_info,
            "vektor": []
        })

        # Take a screenshot of the table
        fig, ax = plt.subplots(figsize=(12, 8))  # Adjust the size as needed
        ax.xaxis.set_visible(False)
        ax.yaxis.set_visible(False)
        ax.set_frame_on(False)
        tbl = table(ax, df, loc='center', cellLoc='center', colWidths=[0.1] * len(df.columns))
        tbl.auto_set_font_size(False)
        tbl.set_fontsize(10)
        tbl.scale(1.2, 1.2)  # Adjust scaling as needed
        plt.savefig(os.path.join(output_dir, f'{filename}_{sheet_name}_table.png'), bbox_inches='tight')
        plt.close(fig)
        print("json")
        table_img_path = os.path.join(output_dir, f'{filename}_{sheet_name}_table.png')
        json_data["data"].append({
            "id": len(json_data["data"]) + 1,
            "titel": f'Table from {filename} sheet {sheet_name}',
            "text": "",
            "vektor": [],
            "path": table_img_path
        })
    """
    except Exception as e:
        print(f"Error processing XLSX {filepath}: {e}")
        return {}
    """
    return json_data


def extract_info(filepath: str, output_dir: str) -> Dict[str, Any]:
    """Extract information from the given file and save it in JSON format."""

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    if filepath.endswith('.pdf'):
        json_data = process_pdf(filepath, output_dir)
    elif filepath.endswith('.docx'):
        json_data = process_docx(filepath, output_dir)
    elif filepath.endswith('.xlsx'):
        json_data = process_xlsx(filepath, output_dir)
    else:
        raise ValueError('Unsupported file type')

    json_path = os.path.join(output_dir, f'{os.path.basename(filepath)}.json')
    with open(json_path, 'w') as json_file:
        json.dump(json_data, json_file, indent=4)
    """
    except Exception as e:
        print(f"Error extracting information from {filepath}: {e}")
        return {}
    """
    return json_data

